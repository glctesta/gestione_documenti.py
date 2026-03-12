# -*- coding: utf-8 -*-
"""
Generazione automatica di rapporti di attività per ospiti di società terze.
Genera documenti Word (Richiesta, Accettazione, Rapporto) e li salva in DB.
"""

import io
import os
import logging
from datetime import datetime, timedelta, date
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger("TraceabilityRS")

# ================================================================
# UTILITY: Calcolo giorni lavorativi
# ================================================================

def _add_working_days(start_date, days):
    """Aggiunge N giorni lavorativi a una data (esclude weekend)."""
    current = start_date
    added = 0
    direction = 1 if days >= 0 else -1
    target = abs(days)
    while added < target:
        current += timedelta(days=direction)
        if current.weekday() < 5:  # Lun-Ven
            added += 1
    return current


def _subtract_working_days(start_date, days):
    """Sottrae N giorni lavorativi da una data."""
    return _add_working_days(start_date, -days)


# ================================================================
# DOCUMENT GENERATOR
# ================================================================

class GuestActivityReportGenerator:
    """Genera i documenti Word per le attività degli ospiti."""

    def __init__(self, db):
        self.db = db
        self._settings_cache = {}

    # --------------------------------------------------------
    # Settings helper
    # --------------------------------------------------------
    def _get_setting(self, atribute):
        """Legge un valore da traceability_rs.dbo.Settings."""
        if atribute in self._settings_cache:
            return self._settings_cache[atribute]
        try:
            cursor = self.db.conn.cursor()
            cursor.execute(
                "SELECT [value] FROM traceability_rs.dbo.Settings WHERE atribute = ?",
                (atribute,))
            row = cursor.fetchone()
            cursor.close()
            val = row[0] if row else ''
            self._settings_cache[atribute] = val
            return val
        except Exception as e:
            logger.error(f"Errore lettura setting '{atribute}': {e}")
            return ''

    def _get_contract_info(self, plan_to_charge_id):
        """Recupera le info contratto per una società."""
        try:
            cursor = self.db.conn.cursor()
            cursor.execute("""
                SELECT ContractNumber, ContractDate, ContractDescription
                FROM Employee.dbo.VisitorContractInfo
                WHERE VisitorPlanToChargeID = ?
            """, (plan_to_charge_id,))
            row = cursor.fetchone()
            cursor.close()
            if row:
                return {
                    'number': row.ContractNumber or '',
                    'date': row.ContractDate,
                    'description': row.ContractDescription or ''
                }
        except Exception as e:
            logger.error(f"Errore lettura contratto: {e}")
        return {'number': '', 'date': None, 'description': ''}

    # --------------------------------------------------------
    # Logo helper
    # --------------------------------------------------------
    def _add_logo(self, doc):
        """Aggiunge il logo Vandewiele in alto al documento."""
        logo_path = os.path.join(os.path.dirname(__file__), 'Logo.png')
        if os.path.exists(logo_path):
            para = doc.add_paragraph()
            run = para.add_run()
            run.add_picture(logo_path, width=Cm(4))
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _set_default_style(self, doc):
        """Imposta lo stile predefinito del documento."""
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    def _add_date_line(self, doc, doc_date):
        """Aggiunge la riga della data."""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.add_run(f"Data: {doc_date.strftime('%d/%m/%Y')}")
        run.font.size = Pt(10)

    def _add_signatories(self, doc, from_name, from_title, to_name, to_title):
        """Aggiunge Da/A con firma."""
        para = doc.add_paragraph()
        run = para.add_run(f"Da: {from_name}")
        run.bold = True
        if from_title:
            para.add_run(f"\n{from_title}")

        para = doc.add_paragraph()
        run = para.add_run(f"A: {to_name}")
        run.bold = True
        if to_title:
            para.add_run(f"\n{to_title}")

        doc.add_paragraph()  # spazio

    # --------------------------------------------------------
    # Doc A - Lettera di Richiesta Intervento
    # --------------------------------------------------------
    def generate_request_letter(self, visitor_data, contract_info, doc_date):
        """
        Genera la Lettera di Richiesta Intervento (Doc A).
        
        Args:
            visitor_data: dict con guest_name, company, start_visit, end_visit, purpose, sponsor
            contract_info: dict con number, date, description
            doc_date: data della lettera
        Returns:
            bytes del documento Word
        """
        doc = Document()
        self._set_default_style(doc)
        self._add_logo(doc)
        self._add_date_line(doc, doc_date)

        chi_richiede = self._get_setting('chi_richiede')
        chi_richiede_titolo = self._get_setting('chi_richiede_titolo')
        chi_invia = self._get_setting('chi_invia')
        chi_invia_titolo = self._get_setting('chi_invia_titolo')

        self._add_signatories(doc, chi_richiede, chi_richiede_titolo,
                              chi_invia, chi_invia_titolo)

        # Oggetto
        para = doc.add_paragraph()
        run = para.add_run("Oggetto: ")
        run.bold = True
        para.add_run(f"Richiesta di intervento — {visitor_data['purpose']}")

        doc.add_paragraph()

        # Riferimento contratto
        contract_ref = ''
        if contract_info['number']:
            contract_ref = f" come previsto dal contratto n. {contract_info['number']}"
            if contract_info['date']:
                contract_ref += f" del {contract_info['date'].strftime('%d/%m/%Y')}"

        # Corpo lettera
        start_str = visitor_data['start_visit'].strftime('%d/%m/%Y')
        end_str = visitor_data['end_visit'].strftime('%d/%m/%Y')

        body = (
            f"Gentile {chi_invia},\n\n"
            f"con la presente, Vi chiediamo cortesemente di mettere a disposizione "
            f"il/la Sig./Sig.ra {visitor_data['guest_name']} "
            f"della società {visitor_data['company']}{contract_ref}, "
            f"per un intervento di supporto tecnico/formazione presso la nostra sede "
            f"di Vandewiele Romania.\n\n"
            f"Periodo richiesto: dal {start_str} al {end_str}\n\n"
            f"Oggetto dell'intervento: {visitor_data['purpose']}\n\n"
            f"La persona di riferimento presso la nostra sede sarà: "
            f"{visitor_data.get('sponsor', 'N/A')}.\n\n"
            f"Restiamo a disposizione per qualsiasi chiarimento.\n\n"
            f"Cordiali saluti,\n\n\n"
            f"{chi_richiede}\n{chi_richiede_titolo}\n"
            f"Vandewiele Romania"
        )

        for line in body.split('\n'):
            doc.add_paragraph(line)

        return self._doc_to_bytes(doc)

    # --------------------------------------------------------
    # Doc B - Lettera di Accettazione
    # --------------------------------------------------------
    def generate_acceptance_letter(self, visitor_data, contract_info, doc_date):
        """
        Genera la Lettera di Accettazione (Doc B).
        """
        doc = Document()
        self._set_default_style(doc)
        self._add_logo(doc)
        self._add_date_line(doc, doc_date)

        chi_richiede = self._get_setting('chi_richiede')
        chi_richiede_titolo = self._get_setting('chi_richiede_titolo')
        chi_invia = self._get_setting('chi_invia')
        chi_invia_titolo = self._get_setting('chi_invia_titolo')

        # Invertito: da chi_invia a chi_richiede
        self._add_signatories(doc, chi_invia, chi_invia_titolo,
                              chi_richiede, chi_richiede_titolo)

        # Oggetto
        para = doc.add_paragraph()
        run = para.add_run("Oggetto: ")
        run.bold = True
        para.add_run(f"Accettazione richiesta di intervento — {visitor_data['purpose']}")

        doc.add_paragraph()

        # Riferimento contratto
        contract_ref = ''
        if contract_info['number']:
            contract_ref = f"in conformità al contratto n. {contract_info['number']}"
            if contract_info['date']:
                contract_ref += f" del {contract_info['date'].strftime('%d/%m/%Y')}"
            contract_ref += ', '

        start_str = visitor_data['start_visit'].strftime('%d/%m/%Y')
        end_str = visitor_data['end_visit'].strftime('%d/%m/%Y')

        body = (
            f"Gentile {chi_richiede},\n\n"
            f"in riscontro alla Vostra richiesta, {contract_ref}"
            f"Vi confermiamo la disponibilità del/della Sig./Sig.ra "
            f"{visitor_data['guest_name']} della società {visitor_data['company']} "
            f"per un intervento presso la Vostra sede.\n\n"
            f"Dettagli dell'intervento:\n"
            f"• Periodo: dal {start_str} al {end_str}\n"
            f"• Oggetto: {visitor_data['purpose']}\n"
            f"• Referente in sede: {visitor_data.get('sponsor', 'N/A')}\n\n"
            f"Il/La Sig./Sig.ra {visitor_data['guest_name']} è a disposizione "
            f"per il periodo indicato e si atterrà alle normative di sicurezza "
            f"vigenti presso la Vostra sede.\n\n"
            f"Cordiali saluti,\n\n\n"
            f"{chi_invia}\n{chi_invia_titolo}\n"
            f"{visitor_data['company']}"
        )

        for line in body.split('\n'):
            doc.add_paragraph(line)

        return self._doc_to_bytes(doc)

    # --------------------------------------------------------
    # Doc C - Rapporto di Attività
    # --------------------------------------------------------
    def generate_activity_report(self, visitor_data, contract_info, doc_date,
                                  activity_description=''):
        """
        Genera il Rapporto di Attività (Doc C).
        """
        doc = Document()
        self._set_default_style(doc)
        self._add_logo(doc)
        self._add_date_line(doc, doc_date)

        chi_richiede = self._get_setting('chi_richiede')
        chi_richiede_titolo = self._get_setting('chi_richiede_titolo')

        # Intestazione
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("RAPPORTO DI ATTIVITÀ")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

        doc.add_paragraph()

        start_str = visitor_data['start_visit'].strftime('%d/%m/%Y')
        end_str = visitor_data['end_visit'].strftime('%d/%m/%Y')

        # Riferimento contratto
        contract_ref = ''
        if contract_info['number']:
            contract_ref = (
                f"come previsto dal contratto n. {contract_info['number']}"
            )
            if contract_info['date']:
                contract_ref += f" del {contract_info['date'].strftime('%d/%m/%Y')}"
            if contract_info['description']:
                contract_ref += f" ({contract_info['description']})"

        # Prefazione
        preface = (
            f"Il presente rapporto documenta le attività svolte dal/dalla "
            f"Sig./Sig.ra {visitor_data['guest_name']} della società "
            f"{visitor_data['company']} presso la sede di Vandewiele Romania "
            f"nel periodo dal {start_str} al {end_str}."
        )
        if contract_ref:
            preface += f"\n\nL'intervento è stato effettuato {contract_ref}."

        doc.add_paragraph(preface)
        doc.add_paragraph()

        # Tabella riepilogativa
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        cells = table.rows[0].cells
        cells[0].text = 'Formatore/Consulente'
        cells[1].text = visitor_data['guest_name']
        cells = table.rows[1].cells
        cells[0].text = 'Società'
        cells[1].text = visitor_data['company']
        cells = table.rows[2].cells
        cells[0].text = 'Periodo'
        cells[1].text = f"{start_str} — {end_str}"
        cells = table.rows[3].cells
        cells[0].text = 'Oggetto'
        cells[1].text = visitor_data['purpose']
        cells = table.rows[4].cells
        cells[0].text = 'Referente VR'
        cells[1].text = visitor_data.get('sponsor', 'N/A')

        doc.add_paragraph()

        # Attività svolte
        para = doc.add_paragraph()
        run = para.add_run("Attività svolte:")
        run.bold = True
        run.font.size = Pt(12)

        if activity_description:
            doc.add_paragraph(activity_description)
        else:
            doc.add_paragraph(
                f"Durante il periodo di permanenza, il/la Sig./Sig.ra "
                f"{visitor_data['guest_name']} ha svolto attività di "
                f"{visitor_data['purpose']}."
            )

        doc.add_paragraph()

        # Firma
        doc.add_paragraph(f"Vandewiele Romania, {doc_date.strftime('%d/%m/%Y')}")
        doc.add_paragraph()
        doc.add_paragraph(f"{chi_richiede}")
        doc.add_paragraph(f"{chi_richiede_titolo}")

        return self._doc_to_bytes(doc)

    # --------------------------------------------------------
    # Helpers
    # --------------------------------------------------------
    def _doc_to_bytes(self, doc):
        """Converte un Document python-docx in bytes."""
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    # --------------------------------------------------------
    # Processo completo per un visitatore
    # --------------------------------------------------------
    def process_visitor(self, visitor_id, activity_description='', created_by='System'):
        """
        Genera tutti i documenti per un visitatore e li salva in DB.
        
        Args:
            visitor_id: ID del visitatore
            activity_description: testo opzionale attività svolte
            created_by: utente che ha lanciato il processo
        Returns:
            VisitorActivityReportId o None
        """
        try:
            # 1. Carica dati visitatore
            cursor = self.db.conn.cursor()
            cursor.execute("""
                SELECT v.VisitorId, v.CompanyName, v.GuestName,
                       v.StartVisit, v.EndVisit, v.Pourpose, v.SponsorGuy,
                       v.VisitorDataId
                FROM Employee.dbo.Visitors v
                WHERE v.VisitorId = ?
            """, (visitor_id,))
            vrow = cursor.fetchone()
            if not vrow:
                logger.error(f"Visitatore {visitor_id} non trovato")
                cursor.close()
                return None

            visitor_data = {
                'guest_name': vrow.GuestName or '',
                'company': vrow.CompanyName or '',
                'start_visit': vrow.StartVisit,
                'end_visit': vrow.EndVisit,
                'purpose': vrow.Pourpose or '',
                'sponsor': vrow.SponsorGuy or '',
                'visitor_data_id': vrow.VisitorDataId
            }

            # 2. Cerca società fatturante (MustCharged=1)
            cursor.execute("""
                SELECT vpc.VisitorPlanToChargeID
                FROM Employee.dbo.VisitorPlanToCharges vpc
                WHERE vpc.CompanyName = ? AND vpc.MustCharged = 1
            """, (visitor_data['company'],))
            charge_row = cursor.fetchone()
            if not charge_row:
                logger.info(f"Società '{visitor_data['company']}' non fatturante, skip")
                cursor.close()
                return None

            plan_to_charge_id = charge_row.VisitorPlanToChargeID

            # 3. Contratto
            contract_info = self._get_contract_info(plan_to_charge_id)

            # 4. Calcola date documenti
            request_date = _subtract_working_days(visitor_data['start_visit'], 5)
            acceptance_date = _add_working_days(request_date, 2)
            report_date = _add_working_days(visitor_data['end_visit'], 1)

            # 5. Genera documenti
            doc_a = self.generate_request_letter(visitor_data, contract_info, request_date)
            doc_b = self.generate_acceptance_letter(visitor_data, contract_info, acceptance_date)
            doc_c = self.generate_activity_report(visitor_data, contract_info, report_date,
                                                   activity_description)

            # 6. Salva in DB
            cursor.execute("""
                INSERT INTO Employee.dbo.VisitorActivityReports
                    (VisitorId, RequestLetterDoc, AcceptanceLetterDoc, ActivityReportDoc,
                     RequestLetterDate, AcceptanceLetterDate, ActivityReportDate,
                     ActivityDescription, CreatedBy)
                OUTPUT INSERTED.VisitorActivityReportId
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (visitor_id, doc_a, doc_b, doc_c,
                  request_date, acceptance_date, report_date,
                  activity_description or None, created_by))

            result = cursor.fetchone()
            report_id = result[0] if result else None
            self.db.conn.commit()
            cursor.close()

            logger.info(f"Documenti generati per visitatore {visitor_id}, "
                        f"ReportId={report_id}")
            return report_id

        except Exception as e:
            logger.error(f"Errore generazione documenti per visitatore {visitor_id}: {e}")
            return None

    # --------------------------------------------------------
    # Invio email rapporto
    # --------------------------------------------------------
    def send_activity_email(self, report_id):
        """
        Invia il rapporto di attività via email.
        TO: chi_richiede_email, chi_invia_email
        CC: email dell'ospite
        """
        from email_connector import EmailSender

        try:
            cursor = self.db.conn.cursor()
            cursor.execute("""
                SELECT var.*, v.GuestName, v.CompanyName, v.Pourpose,
                       v.StartVisit, v.EndVisit, v.VisitorDataId
                FROM Employee.dbo.VisitorActivityReports var
                INNER JOIN Employee.dbo.Visitors v ON var.VisitorId = v.VisitorId
                WHERE var.VisitorActivityReportId = ?
            """, (report_id,))
            row = cursor.fetchone()
            if not row:
                logger.error(f"Report {report_id} non trovato")
                cursor.close()
                return False

            # Email destinatari
            to_richiede = self._get_setting('chi_richiede_email')
            to_invia = self._get_setting('chi_invia_email')
            to_list = [e.strip() for e in [to_richiede, to_invia] if e.strip()]

            if not to_list:
                logger.error("Nessun destinatario email configurato")
                cursor.close()
                return False

            # CC: email ospite
            cc_email = None
            if row.VisitorDataId:
                cursor.execute("""
                    SELECT EmailAddress FROM Employee.dbo.VisitorData
                    WHERE VisitorDataID = ?
                """, (row.VisitorDataId,))
                vd_row = cursor.fetchone()
                if vd_row and vd_row.EmailAddress:
                    cc_email = vd_row.EmailAddress.strip()

            # Prepara allegati Word
            import tempfile
            temp_files = []
            attachments = []

            logo_path = os.path.join(os.path.dirname(__file__), 'Logo.png')

            for doc_name, doc_bytes in [
                ('Richiesta_Intervento.docx', row.RequestLetterDoc),
                ('Accettazione.docx', row.AcceptanceLetterDoc),
                ('Rapporto_Attivita.docx', row.ActivityReportDoc)
            ]:
                if doc_bytes:
                    tmp = tempfile.NamedTemporaryFile(
                        suffix='.docx', prefix=doc_name.replace('.docx', '_'),
                        delete=False)
                    tmp.write(doc_bytes)
                    tmp.close()
                    temp_files.append(tmp.name)
                    attachments.append(tmp.name)

            if os.path.exists(logo_path):
                attachments.insert(0, ('inline', logo_path, 'company_logo'))

            start_str = row.StartVisit.strftime('%d/%m/%Y') if row.StartVisit else ''
            end_str = row.EndVisit.strftime('%d/%m/%Y') if row.EndVisit else ''

            body_html = f"""
            <html>
            <body style="font-family: Arial, sans-serif; font-size: 13px; color: #333;">
                <img src="cid:company_logo" alt="Vandewiele" 
                     style="width: 160px; margin-bottom: 15px;" /><br/>

                <h2 style="color: #1565C0;">Rapporto di Attività — {row.GuestName}</h2>

                <p>In allegato la documentazione relativa all'intervento di:</p>

                <table style="border-collapse: collapse; margin: 15px 0; font-size: 13px;">
                    <tr style="background-color: #E3F2FD;">
                        <td style="padding: 8px 15px; font-weight: bold;">Formatore</td>
                        <td style="padding: 8px 15px;">{row.GuestName}</td>
                    </tr>
                    <tr>
                        <td style="padding: 8px 15px; font-weight: bold;">Società</td>
                        <td style="padding: 8px 15px;">{row.CompanyName}</td>
                    </tr>
                    <tr style="background-color: #E3F2FD;">
                        <td style="padding: 8px 15px; font-weight: bold;">Periodo</td>
                        <td style="padding: 8px 15px;">{start_str} — {end_str}</td>
                    </tr>
                    <tr>
                        <td style="padding: 8px 15px; font-weight: bold;">Oggetto</td>
                        <td style="padding: 8px 15px;">{row.Pourpose}</td>
                    </tr>
                </table>

                <p>Documenti allegati:</p>
                <ul>
                    <li>Lettera di Richiesta Intervento</li>
                    <li>Lettera di Accettazione</li>
                    <li>Rapporto di Attività</li>
                </ul>

                <hr style="border: 1px solid #ddd; margin-top: 20px;"/>
                <p style="color: #999; font-size: 10px;">
                    Messaggio automatico — TraceabilityRS</p>
            </body>
            </html>
            """

            sender = EmailSender()
            sender.send_email(
                to_email='; '.join(to_list),
                subject=f"Rapporto Attività — {row.GuestName} — {row.CompanyName}",
                body=body_html,
                is_html=True,
                attachments=attachments if attachments else None,
                cc_emails=cc_email
            )

            # Aggiorna DB
            cursor.execute("""
                UPDATE Employee.dbo.VisitorActivityReports
                SET EmailSentDate = GETDATE(), EmailSentTo = ?
                WHERE VisitorActivityReportId = ?
            """, ('; '.join(to_list), report_id))
            self.db.conn.commit()
            cursor.close()

            # Cleanup temp
            for f in temp_files:
                try:
                    os.unlink(f)
                except Exception:
                    pass

            logger.info(f"Email rapporto inviata: report={report_id}, to={to_list}")
            return True

        except Exception as e:
            logger.error(f"Errore invio email rapporto {report_id}: {e}")
            return False

    # --------------------------------------------------------
    # Batch: processa visitatori partiti in un intervallo
    # --------------------------------------------------------
    def process_departed_visitors(self, from_date=None, to_date=None):
        """
        Cerca visitatori partiti nell'intervallo [from_date, to_date]
        con società MustCharged=1 e genera i documenti se non già generati.

        Default: dal 1° gennaio dell'anno corrente a ieri.
        """
        try:
            if to_date is None:
                to_date = date.today() - timedelta(days=1)
            if from_date is None:
                from_date = date(to_date.year, 1, 1)

            logger.info(f"Batch rapporti attività: periodo {from_date} → {to_date}")

            cursor = self.db.conn.cursor()
            cursor.execute("""
                SELECT v.VisitorId, v.CompanyName, v.GuestName
                FROM Employee.dbo.Visitors v
                INNER JOIN Employee.dbo.VisitorPlanToCharges vpc
                    ON v.CompanyName = vpc.CompanyName
                WHERE CAST(v.EndVisit AS DATE) BETWEEN ? AND ?
                  AND vpc.MustCharged = 1
                  AND v.IsActive = 1
                  AND NOT EXISTS (
                      SELECT 1 FROM Employee.dbo.VisitorActivityReports var
                      WHERE var.VisitorId = v.VisitorId
                  )
            """, (from_date, to_date))

            visitors = cursor.fetchall()
            cursor.close()

            logger.info(f"Batch: trovati {len(visitors)} visitatori senza report")

            count = 0
            for v in visitors:
                try:
                    report_id = self.process_visitor(v.VisitorId, created_by='AutoBatch')
                    if report_id:
                        self.send_activity_email(report_id)
                        count += 1
                        logger.info(f"Batch: report generato per {v.GuestName} ({v.CompanyName})")
                except Exception as ve:
                    logger.error(f"Batch: errore per VisitorId={v.VisitorId}: {ve}")

            logger.info(f"Batch rapporti attività completato: {count}/{len(visitors)} generati")
            return count

        except Exception as e:
            logger.error(f"Errore batch rapporti attività: {e}")
            return 0

